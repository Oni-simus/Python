import os
import datetime
import openpyxl
import docx
import tkinter as tk
import time
import progressbar
import re


print('Developed by Jayric Maning')
print('TaskBot version 1.2.0 Beta')
print('')
print('reminders!')
print('1.Do make sure that your xlsx file is stored/placed at your desktop\n or screen,no folders!')
os.chdir('c:\\Users\\user\\Desktop')
pbar=progressbar.ProgressBar()
wb = openpyxl.load_workbook('MI DESTINO 3 - HRLD.xlsx')
print('harold,nadine,gene,aubrey,jayric,elija,MARK')
writer_name=input('Who are you?')
sheet = wb.get_sheet_by_name(writer_name)
maxrow = sheet.max_row
maxcol = sheet.max_column


#this section is for all the interchangable variables to be used in automating
#the creation of documents
essay_type=[]
titles=[]
writer=[]
doc_number=[0,1,0,2,0,3,0,4,0,5,0]
link=[]
keyword=[]
for i in range(maxrow-4,maxrow+1,1):
    essay_type+=(i,sheet.cell(row=i,column=2).value)
for i in range(maxrow-4,maxrow+1,1):
    titles+=(i,sheet.cell(row=i,column=3).value)
for i in range(maxrow-4,maxrow+1,1):
    writer+=(i,sheet.cell(row=i,column=4).value)
for i in range(maxrow-4,maxrow+1,1):
    link+=(i,sheet.cell(row=i,column=6).value)
for i in range(maxrow-4,maxrow+1,1):
    keyword+=(i,sheet.cell(row=i,column=1).value)

    

#this section will concantenate the string into integer
#then add 1, then concantenate it again to string
day =datetime.datetime.today().strftime('%d')
Integer =(int(day)+1)
String= (str(Integer))
FolderName = datetime.datetime.today().strftime('%m-'+String+'-%Y')

#this section is for creating the folder 
os.chdir("C:\\Users\\user\\Desktop\\")
os.mkdir("" + FolderName)
doc = docx.Document()
os.chdir("C:\\Users\\user\\Desktop\\" +FolderName)


#this section will make a loop to create all of the documents with keyword/title
#and link.


print('')
print('')
print('')

class App(tk.Frame):
     def __init__(self,object):
          tk.Frame.__init__(self,object)
          self.pack()
          self.master.title('Task Bot')
          tk.Label(self, text= 'What type of set are you taking?').pack()
          tk.Button(self,text = 'MATT',command=self.Matt).pack(side='left')
          tk.Button(self,text = 'REAL',command=self.Real).pack(side='left')
          tk.Button(self,text = 'DOLLAR',command=self.Dollar).pack(side='left')
          tk.Button(self,text = 'JUNEJA',command=self.Juneja).pack(side='left')
          tk.Button(self,text = 'FIFTY',command=self.Fifty).pack(side='left')
          tk.Button(self,text='cancel',command=self.click_cancel).pack(side='right')


     def Matt(self):
         i=1
         while i < 10:
             doc = docx.Document()
             doc.add_paragraph('Keyword:'+str(titles[i]));
             doc.add_paragraph('Research Link:'+str(link[i]));
             doc.add_paragraph('Target Site:'+str(keyword[i]));
             try:
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             except:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             finally:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             i+=2

     def Dollar(self):
         i=1
         while i < 10:
             doc = docx.Document()
             doc.add_paragraph('Keyword:'+str(titles[i]));
             doc.add_paragraph('Research Link:'+str(link[i]));
             doc.add_paragraph('Target Site:'+str(keyword[i]));
             try:
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             except:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             finally:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');

             i+=2

                
     def Real(self):
         i=1
         while i < 10:
             doc = docx.Document()
             doc.add_paragraph('Title:'+str(titles[i]));
             doc.add_paragraph('Research Link:'+str(link[i]));
             doc.add_paragraph('Keywords:'+str(keyword[i]));
             try:
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             except:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             finally:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');

             i+=2

             
     def Juneja(self):
         i=1
         while i < 10:
             doc = docx.Document()
             doc.add_paragraph('Keyword:'+str(titles[i]));
             doc.add_paragraph('Research Link:'+str(link[i]));
             doc.add_paragraph('Replacement Link:'+str(keyword[i]));
             try:
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             except:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');
             finally:
                 titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
                 doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                      +str(titles[i])+' '+str(writer[i])+'.docx');

             i+=2
             
                
     def Fifty(self):
        i=1
        doc.add_paragraph('Keyword:'+str(titles[i]));
        doc.add_paragraph('Research Link:'+str(link[i]));
        doc.add_paragraph('Target Site:'+str(keyword[i]));
        try:
            doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                +str(titles[i])+' '+str(writer[i])+'.docx');
        except:
            titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
            doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                +str(titles[i])+' '+str(writer[i])+'.docx');
        finally:
            titles[i] = titles[i].replace(':','').replace('?','').replace('*','')
            doc.save(str(doc_number[i])+str(essay_type[i])+' '\
                +str(titles[i])+' '+str(writer[i])+'.docx');


        i+=2

                
     def click_cancel(self):
         print('good bye')
         quit()


if __name__ =='__main__':
    root = tk.Tk()
    app = App(root)
    app.mainloop()



